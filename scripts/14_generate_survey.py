"""
14_generate_survey.py
=====================
Generate blind expert survey comparing SciBERT vs TOD-SciBERT.
Experts don't know which is A or B.

Output: E:\TOD_DAPT\survey\expert_survey.docx
        E:\TOD_DAPT\survey\answer_key.txt  (which A/B is which model)
"""

import os, re, random, torch
import numpy as np
import fitz
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import AutoTokenizer, AutoModel
from nltk.tokenize import sent_tokenize
import nltk

try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab', quiet=True)

# ── Paths ──────────────────────────────────────────────────
SCIBERT_PATH     = "allenai/scibert_scivocab_uncased"
TOD_SCIBERT_PATH = r"E:\TOD_DAPT\model\tod_scibert"
CORPUS_PATH      = r"E:\TOD_DAPT\corpus\dapt_corpus.txt"
OUTPUT_DIR       = r"E:\TOD_DAPT\survey"
os.makedirs(OUTPUT_DIR, exist_ok=True)

TOP_K = 5  # sentences per set

# ── 15 Expert Queries ──────────────────────────────────────
QUERIES = [
    ("Q1",  "Walkability",           "Walkability is a key factor in achieving sustainable transit-oriented development."),
    ("Q2",  "Transit Access",        "High transit accessibility improves the sustainability outcomes of TOD areas."),
    ("Q3",  "Land Use Mix",          "Mixed land use within station catchment zones strengthens transit-oriented development."),
    ("Q4",  "Density",               "Higher residential and commercial density near transit stations supports TOD sustainability."),
    ("Q5",  "Parking Reduction",     "Reducing surface parking in TOD zones leads to more sustainable urban outcomes."),
    ("Q6",  "Cycling Infrastructure","Dedicated cycling infrastructure is essential for sustainable TOD zones."),
    ("Q7",  "Green Space",           "Integrating green spaces and parks within TOD areas improves sustainability performance."),
    ("Q8",  "Affordable Housing",    "Affordable housing provision is a critical component of sustainable TOD."),
    ("Q9",  "Station Area Design",   "Well-designed station areas positively influence TOD sustainability outcomes."),
    ("Q10", "Urban Design Quality",  "High-quality urban design in TOD zones contributes to sustainable development."),
    ("Q11", "Carbon Emission",       "Transit-oriented development reduces per-capita carbon emissions in urban areas."),
    ("Q12", "Social Equity",         "Social equity considerations are fundamental to sustainable TOD planning."),
    ("Q13", "Connectivity",          "Strong street connectivity within TOD zones supports sustainable transit use."),
    ("Q14", "Energy Efficiency",     "Energy-efficient building practices in TOD areas enhance sustainability outcomes."),
    ("Q15", "Economic Viability",    "Economic vitality of station areas is necessary for long-term TOD sustainability."),
]


# ── Helpers ────────────────────────────────────────────────
def mean_pool(token_emb, attn_mask):
    mask = attn_mask.unsqueeze(-1).expand(token_emb.size()).float()
    return (token_emb * mask).sum(1) / mask.sum(1).clamp(min=1e-9)


def embed(model, tokenizer, texts, device, batch_size=32):
    all_embs = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i+batch_size]
        enc = tokenizer(batch, padding=True, truncation=True,
                        max_length=256, return_tensors='pt').to(device)
        with torch.no_grad():
            out = model(**enc)
        emb = mean_pool(out.last_hidden_state, enc['attention_mask'])
        emb = torch.nn.functional.normalize(emb, p=2, dim=1)
        all_embs.append(emb.cpu().numpy())
    return np.vstack(all_embs)


def get_top_sentences(model, tokenizer, query, sentences, sent_embs, device, k=5):
    q_emb = embed(model, tokenizer, [query], device)[0]
    scores = sent_embs @ q_emb
    top_idx = np.argsort(scores)[::-1][:k]
    return [sentences[i] for i in top_idx]


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    return p


# ── Main ───────────────────────────────────────────────────
def main():
    device = "cuda" if torch.cuda.is_available() else "cpu"
    print("=" * 55)
    print("  Expert Survey Generator")
    print("=" * 55)
    print(f"\n  Device: {device}")

    # Load corpus
    print("\n  Loading corpus...")
    with open(CORPUS_PATH, 'r', encoding='utf-8') as f:
        sentences = [l.strip() for l in f if len(l.strip()) > 40]
    sentences = sentences[:30000]  # use first 30K for speed
    print(f"  Sentences: {len(sentences):,}")

    # Load models
    print("\n  Loading SciBERT...")
    sci_tok = AutoTokenizer.from_pretrained(SCIBERT_PATH)
    sci_mod = AutoModel.from_pretrained(SCIBERT_PATH).to(device).eval()

    print("  Loading TOD-SciBERT...")
    tod_tok = AutoTokenizer.from_pretrained(TOD_SCIBERT_PATH)
    tod_mod = AutoModel.from_pretrained(TOD_SCIBERT_PATH).to(device).eval()

    # Pre-embed all sentences
    print("\n  Embedding corpus sentences (this takes ~2 min)...")
    sci_embs = embed(sci_mod, sci_tok, sentences, device, batch_size=64)
    tod_embs = embed(tod_mod, tod_tok, sentences, device, batch_size=64)
    print("  Done!")

    # Retrieve sentences for each query
    print("\n  Retrieving top sentences per query...")
    query_results = []
    answer_key    = []

    for qid, aspect, query in QUERIES:
        sci_sents = get_top_sentences(sci_mod, sci_tok, query, sentences, sci_embs, device)
        tod_sents = get_top_sentences(tod_mod, tod_tok, query, sentences, tod_embs, device)

        # Randomly assign A/B
        if random.random() > 0.5:
            set_a, set_b = sci_sents, tod_sents
            key = f"{qid}: A=SciBERT, B=TOD-SciBERT"
        else:
            set_a, set_b = tod_sents, sci_sents
            key = f"{qid}: A=TOD-SciBERT, B=SciBERT"

        query_results.append((qid, aspect, query, set_a, set_b))
        answer_key.append(key)
        print(f"  {qid} ✓")

    # ── Build Word Document ───────────────────────────────────
    print("\n  Building Word document...")
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Title
    title = doc.add_heading("Expert Survey: TOD Language Model Evaluation", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)

    doc.add_paragraph()

    # Introduction box
    intro = doc.add_paragraph()
    intro.add_run("Purpose of this Survey\n").bold = True
    intro.add_run(
        "This survey evaluates whether a domain-adapted language model "
        "(trained specifically on Transit-Oriented Development research papers) "
        "retrieves more relevant evidence than a general scientific language model.\n\n"
        "You will be shown 15 TOD-related statements. For each statement, "
        "two sets of retrieved sentences (Set A and Set B) are presented — each retrieved "
        "by a different language model. Please evaluate which set of sentences better "
        "supports or is more relevant to the given statement."
    )
    intro.paragraph_format.left_indent  = Cm(0.5)
    intro.paragraph_format.right_indent = Cm(0.5)

    doc.add_paragraph()

    # Expert info
    add_para(doc, "Expert Information", bold=True, size=12)
    for field in ["Name:", "Designation:", "Institution:", "Years of TOD/Urban Planning Experience:"]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(field + "  ___________________________________")

    doc.add_paragraph()

    # Rating instructions
    add_para(doc, "Rating Scale", bold=True, size=12)
    scale = doc.add_paragraph()
    scale.add_run(
        "For each statement, mark ONE option based on which retrieved set better supports the statement:\n"
        "  ⬜  Set A is much more relevant to the statement\n"
        "  ⬜  Set A is slightly more relevant to the statement\n"
        "  ⬜  Both sets are equally relevant\n"
        "  ⬜  Set B is slightly more relevant to the statement\n"
        "  ⬜  Set B is much more relevant to the statement"
    )

    doc.add_page_break()

    # Questions
    for qid, aspect, query, set_a, set_b in query_results:
        # Question header
        h = doc.add_heading(f"{qid}. {aspect}", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)

        # Statement
        p = doc.add_paragraph()
        p.add_run("Statement:  ").bold = True
        p.add_run(query).italic = True

        doc.add_paragraph()

        # Set A
        p = doc.add_paragraph()
        p.add_run("Set A — Retrieved Sentences:").bold = True
        p.add_run(" (from Language Model A)")
        for i, sent in enumerate(set_a, 1):
            p2 = doc.add_paragraph(style='List Number')
            p2.add_run(sent[:300] + ("..." if len(sent) > 300 else ""))
            p2.paragraph_format.left_indent = Cm(0.5)

        doc.add_paragraph()

        # Set B
        p = doc.add_paragraph()
        p.add_run("Set B — Retrieved Sentences:").bold = True
        p.add_run(" (from Language Model B)")
        for i, sent in enumerate(set_b, 1):
            p2 = doc.add_paragraph(style='List Number')
            p2.add_run(sent[:300] + ("..." if len(sent) > 300 else ""))
            p2.paragraph_format.left_indent = Cm(0.5)

        doc.add_paragraph()

        # Rating
        p = doc.add_paragraph()
        p.add_run("Your Rating:").bold = True
        rating_text = doc.add_paragraph(
            "  ⬜ Set A is much more relevant     "
            "⬜ Set A is slightly more relevant     "
            "⬜ Equal     "
            "⬜ Set B is slightly more relevant     "
            "⬜ Set B is much more relevant"
        )

        # Comments
        doc.add_paragraph("Comments (optional): ____________________________________________")

        doc.add_paragraph()
        doc.add_paragraph("─" * 80)
        doc.add_paragraph()

    # Thank you
    thanks = doc.add_paragraph()
    thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    thanks.add_run(
        "Thank you for your valuable time and expertise!\n"
        "Please return the completed survey to: [your email]"
    ).italic = True

    # Save document
    survey_path = os.path.join(OUTPUT_DIR, 'expert_survey.docx')
    doc.save(survey_path)

    # Save answer key
    key_path = os.path.join(OUTPUT_DIR, 'answer_key.txt')
    with open(key_path, 'w') as f:
        f.write("ANSWER KEY (Confidential — Do not share with experts)\n")
        f.write("=" * 50 + "\n\n")
        for k in answer_key:
            f.write(k + "\n")

    print(f"\n✅ Survey saved: {survey_path}")
    print(f"✅ Answer key : {key_path}")
    print(f"\n  Experts-কে দাও: expert_survey.docx")
    print(f"  নিজে রাখো:      answer_key.txt")


if __name__ == "__main__":
    random.seed(42)
    main()
